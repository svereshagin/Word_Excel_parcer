from docx import Document
import os
from excel_reader import data
input_file = 'Отчет_шаблон.docx'
old_name = 'Р.С. Абрамов'


def replace_name_in_docx(input_file, old_name, new_name, output_file):
    doc = Document(input_file)

    for para in doc.paragraphs:
        if old_name in para.text:
            para.text = para.text.replace(old_name, new_name)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_name in cell.text:
                    cell.text = cell.text.replace(old_name, new_name)


    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    doc.save(output_file)
    print(f"Документ сохранен как {output_file}")


for i in data:
    new_name = i[0]
    output_file = os.path.join(f'{i[0]}', f'{i[0]}.docx')
    replace_name_in_docx(input_file, old_name, new_name, output_file)